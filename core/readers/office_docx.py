"""Word documents -> core.ir.Document, via python-docx.

A .docx is a flow format: nothing in the file says where page 4 ends. Word
decides that at layout time from the page size, the fonts installed and the
printer driver, which is why the same file paginates differently on two
machines. So pages here are synthetic and disclosed, exactly as for HTML --
what this reader must NOT do is invent a page count and present it as the
document's own.

The structure that IS declared is worth a great deal: paragraph styles name
their own headings ("Heading 1"), and `w:tbl` is a real table. Both come back
at `native` basis, so an outline of a Word document is stated rather than
guessed from type size the way a PDF's has to be.

Body order matters. python-docx exposes `document.paragraphs` and
`document.tables` as two separate lists, and reading them in that order puts
every table at the end of the document regardless of where it sits. This walks
the body XML instead, so a table between two paragraphs lands between them.
"""

from __future__ import annotations

from core.ir import Block, Document, Span
from core.readers import _flow

HEADING_SIZES = {1: 26.0, 2: 22.0, 3: 18.0, 4: 16.0, 5: 14.0, 6: 13.0}
BODY_SIZE = 12.0
MAX_HEADING_LEVEL = 6

# The OOXML namespace, needed to tell a w:p from a w:tbl while walking the body.
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def open_document(path: str, password: str = "") -> Document:
    """Read a .docx whole. Refuses one too large to hold in memory."""
    import docx

    src = _flow.guard_size(path)
    try:
        handle = docx.Document(str(src))
    except Exception as exc:
        from core.readers import ReaderError

        message = str(exc)
        if "encrypted" in message.lower() or "not a zip" in message.lower():
            raise ReaderError(
                f"{src.name} could not be opened; it may be password-protected or not a .docx.",
                "Word's own encryption is not readable here. Save an unprotected copy first.",
            ) from exc
        raise ReaderError(
            f"{src.name} could not be read as a Word document: {message}",
            "Check the file opens in Word. A .doc (not .docx) needs convert(to='pdf') first.",
        ) from exc

    blocks = list(_walk(handle))
    core = handle.core_properties
    meta = {
        "title": core.title or "",
        "author": core.author or "",
        "headings": sum(1 for b in blocks if b.kind == "heading"),
        "tables": sum(1 for b in blocks if b.kind == "table"),
    }
    doc = _flow.build(str(src), "docx", blocks, meta)
    return doc


def _walk(handle):
    """Yield blocks in body order, paragraphs and tables interleaved."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = handle.element.body
    for child in body.iterchildren():
        if child.tag == f"{W}p":
            block = _paragraph(Paragraph(child, handle))
            if block is not None:
                yield block
        elif child.tag == f"{W}tbl":
            yield _table(Table(child, handle))


def _paragraph(paragraph) -> Block | None:
    text = paragraph.text.strip()
    if not text:
        return None

    level = _heading_level(paragraph)
    if level:
        return _flow.block(text, "heading", "native", HEADING_SIZES[level])
    style = (paragraph.style.name or "").lower()
    if "list" in style or "bullet" in style:
        return _flow.block(text, "list", "native", BODY_SIZE)
    if "caption" in style:
        return _flow.block(text, "caption", "native", BODY_SIZE)
    return _flow.block(text, "para", "native", BODY_SIZE)


def _heading_level(paragraph) -> int:
    """The declared heading level, or 0.

    Read from the STYLE NAME, which is what Word writes and what survives a
    round trip, rather than from the run's font size. Size would find headings
    a writer made by hand-enlarging body text -- but it would also call every
    large word a heading, which is the low-confidence path the PDF reader is
    stuck with and this format does not need.

    "Title" is level 1: it is the document's own name for its top heading, and
    dropping it would leave an outline whose first entry is a section.
    """
    name = (paragraph.style.name or "").strip()
    if name.lower() == "title":
        return 1
    if not name.lower().startswith("heading"):
        return 0
    tail = name[len("heading") :].strip()
    if tail.isdigit():
        return min(int(tail), MAX_HEADING_LEVEL)
    return 0


def _table(table) -> Block:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(cell for cell in row)]
    flat = "\n".join("\t".join(cell for cell in row) for row in rows)
    return Block(kind="table", spans=[Span(text=flat, size=BODY_SIZE)], bbox=None, basis="native", rows=rows)


def page_tables(doc: Document, numbers: list[int]) -> list[dict]:
    from core.readers.html import page_tables as _shared

    return _shared(doc, numbers)


def bookmarks(doc: Document) -> list[dict]:
    """The outline the document's own styles declare."""
    levels = {size: level for level, size in HEADING_SIZES.items()}
    out: list[dict] = []
    for number in sorted(doc.pages):
        for block in doc.pages[number].blocks:
            if block.kind != "heading":
                continue
            size = block.spans[0].size if block.spans else BODY_SIZE
            out.append(
                {
                    "level": levels.get(size, MAX_HEADING_LEVEL),
                    "title": block.text.strip(),
                    "page": number,
                    "basis": "native",
                }
            )
    return out


def probe_extras(doc: Document) -> dict:
    return {key: doc.meta[key] for key in ("title", "author", "headings", "tables") if key in doc.meta}


close_document = _flow.close_document
load_page = _flow.load_page
load_page_words = _flow.load_page
